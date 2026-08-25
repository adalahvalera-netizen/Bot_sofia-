import os
import urllib.request
import json
import telebot
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import openpyxl
from gtts import gTTS
import yt_dlp

# --- CONFIGURACIÓN ---
TELEGRAM_TOKEN = "8993633836:AAGJJHm9_3bSksfglYXs_T_vveLU8ny1h9I"
COHERE_API_KEY = os.getenv("COHERE_API_KEY", "")

bot = telebot.TeleBot(TELEGRAM_TOKEN)

# --- IA COHERE ---
def consultar_ia_gratis(prompt_usuario):
    if not COHERE_API_KEY:
        return "Falta agregar la variable COHERE_API_KEY en Railway."
    
    try:
        url = "https://api.cohere.com/v1/chat"
        instrucciones = (
            "Eres Sofía, una asistente virtual amigable creada y desarrollada por Abdallah. "
            "Responde de forma breve, útil y natural en español."
        )
        
        payload = json.dumps({
            "message": prompt_usuario,
            "preamble": instrucciones
        }).encode("utf-8")

        req = urllib.request.Request(
            url, 
            data=payload, 
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {COHERE_API_KEY}"
            },
            method="POST"
        )

        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            return data.get("text", "").strip()

    except Exception as e:
        print(f"Error Cohere: {e}")
        return "Error al consultar la IA. Revisa tu COHERE_API_KEY."

# --- GENERADORES DE ARCHIVOS ---
def generar_pdf(nombre_archivo, titulo, contenido):
    doc = SimpleDocTemplate(nombre_archivo, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph(titulo, styles['Heading1']),
        Spacer(1, 12),
        Paragraph(contenido, styles['BodyText'])
    ]
    doc.build(story)

def generar_word_texto(nombre_archivo, titulo, texto):
    doc = Document()
    doc.add_heading(titulo, level=1)
    doc.add_paragraph(texto)
    doc.save(nombre_archivo)

def generar_excel(nombre_archivo):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Area", "Estado", "Creado Por"])
    ws.append(["Educación Física", "Activo", "Abdallah"])
    wb.save(nombre_archivo)

# --- DESCARGAR MÚSICA ---
def descargar_musica(nombre_cancion):
    archivo_salida = "cancion.mp3"
    if os.path.exists(archivo_salida):
        os.remove(archivo_salida)

    ydl_opts = {
        'format': 'bestaudio/best',
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'mp3',
            'preferredquality': '192',
        }],
        'outtmpl': 'cancion',
        'default_search': 'ytsearch1:',
        'quiet': True
    }

    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([nombre_cancion])
    
    return archivo_salida

# --- COMANDOS Y MANEJADORES ---
@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    bot.reply_to(message, "¡Hola! Soy Sofía, creada por Abdallah. Puedo generar documentos, buscar canciones, crear imágenes, extraer audios de videos y hablar contigo por voz.")

# 1. Recibir Fotos
@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    bot.send_chat_action(message.chat.id, 'typing')
    bot.reply_to(message, "¡Recibí tu imagen! Qué gran foto me has mandado.")

# 2. Recibir Videos (Extraer Audio)
@bot.message_handler(content_types=['video'])
def handle_video(message):
    bot.send_message(message.chat.id, "🎥 Recibí tu video. Extrayendo el audio, dame un momento...")
    bot.send_chat_action(message.chat.id, 'upload_document')
    
    try:
        file_info = bot.get_file(message.video.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        
        video_path = "video_recibido.mp4"
        audio_path = "audio_extraido.mp3"
        
        with open(video_path, 'wb') as new_file:
            new_file.write(downloaded_file)
            
        os.system(f"ffmpeg -i {video_path} -q:a 0 -map a {audio_path} -y")
        
        if os.path.exists(audio_path):
            with open(audio_path, 'rb') as audio_file:
                bot.send_audio(message.chat.id, audio_file, title="Audio extraído", performer="Sofía Bot")
        else:
            bot.reply_to(message, "No pude extraer el audio de este formato de video.")
            
        for f in [video_path, audio_path]:
            if os.path.exists(f):
                os.remove(f)
                
    except Exception as e:
        print(f"Error procesando video: {e}")
        bot.reply_to(message, "Ocurrió un error al procesar el video.")

# 3. Mensajes de Texto (Comandos, descargas, imágenes y voz)
@bot.message_handler(func=lambda message: True)
def handle_conversation(message):
    user_text = message.text.lower()
    
    # Crear PDF
    if "pdf" in user_text:
        bot.send_chat_action(message.chat.id, 'upload_document')
        archivo = "documento.pdf"
        generar_pdf(archivo, "Documento PDF", "Este es un archivo PDF generado automáticamente por Sofía.")
        with open(archivo, "rb") as f:
            bot.send_document(message.chat.id, f)
        return

    # Crear Word
    if "word" in user_text or "doc" in user_text or "crucigrama" in user_text:
        bot.send_chat_action(message.chat.id, 'upload_document')
        archivo = "documento.docx"
        generar_word_texto(archivo, "Documento de Word", "Este es un documento de Word generado automáticamente por Sofía.")
        with open(archivo, "rb") as f:
            bot.send_document(message.chat.id, f)
        return
        
    # Crear Excel
    if "excel" in user_text or "tabla" in user_text:
        bot.send_chat_action(message.chat.id, 'upload_document')
        archivo = "tabla.xlsx"
        generar_excel(archivo)
        with open(archivo, "rb") as f:
            bot.send_document(message.chat.id, f)
        return

    # Descargar Música (Ej: "descarga coqueta de heredero")
    if "descarga" in user_text or "cancion" in user_text or "busca la canción" in user_text:
        bot.send_message(message.chat.id, "🔍 Buscando y descargando tu música...")
        bot.send_chat_action(message.chat.id, 'upload_document')
        
        try:
            busqueda = user_text.replace("descarga", "").replace("busca la canción", "").replace("cancion", "").strip()
            archivo_audio = descargar_musica(busqueda)
            
            with open(archivo_audio, "rb") as audio:
                bot.send_audio(message.chat.id, audio, title=busqueda.capitalize(), performer="Sofía Bot")
            
            if os.path.exists(archivo_audio):
                os.remove(archivo_audio)
            return
        except Exception as e:
            print(f"Error al descargar: {e}")
            bot.reply_to(message, "No pude descargar la canción. Intenta escribir el nombre exacto.")
            return

    # Crear Imágenes (Ej: "dibuja un gato")
    if "dibuja" in user_text or "crea una imagen" in user_text:
        bot.send_message(message.chat.id, "🎨 Creando tu imagen...")
        bot.send_chat_action(message.chat.id, 'upload_photo')
        
        try:
            prompt = user_text.replace("crea una imagen de", "").replace("crea una imagen", "").replace("dibuja", "").strip()
            prompt_url = prompt.replace(" ", "%20")
            url_imagen = f"https://image.pollinations.ai/prompt/{prompt_url}?width=1024&height=1024&nologo=true"
            bot.send_photo(message.chat.id, url_imagen, caption=f"Aquí tienes: {prompt.capitalize()}")
            return
        except Exception as e:
            print(f"Error al crear imagen: {e}")
            bot.reply_to(message, "No pude generar esa imagen, intenta con otra descripción.")
            return

    # Respuesta por IA en NOTA DE VOZ (Conversación normal)
    bot.send_chat_action(message.chat.id, 'record_audio')
    respuesta_texto = consultar_ia_gratis(message.text)
    
    archivo_voz = "voz_sofia.mp3"
    tts = gTTS(text=respuesta_texto, lang='es', tld='com')
    tts.save(archivo_voz)
    
    with open(archivo_voz, "rb") as voice:
        bot.send_voice(message.chat.id, voice)
        
    if os.path.exists(archivo_voz):
        os.remove(archivo_voz)

if __name__ == "__main__":
    bot.infinity_polling()
    
