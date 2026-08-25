import os
import urllib.request
import json
import re
import telebot
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import openpyxl
from gtts import gTTS
import yt_dlp
import cohere

# --- CONFIGURACIÓN ---
TELEGRAM_TOKEN = "8993633836:AAGJJHm9_3bSksfglYXs_T_vveLU8ny1h9I"
COHERE_API_KEY = os.getenv("COHERE_API_KEY", "").strip()

bot = telebot.TeleBot(TELEGRAM_TOKEN)

# Inicializar cliente de Cohere
co = cohere.Client(COHERE_API_KEY) if COHERE_API_KEY else None

# Diccionario para recordar el modo de cada usuario
modo_usuario = {}

# --- IA COHERE (SDK OFICIAL) ---
def consultar_ia_gratis(prompt_usuario):
    if not COHERE_API_KEY or not co:
        return "Falta agregar la variable COHERE_API_KEY en Railway."
    
    try:
        instrucciones = (
            "Eres Sofía, una asistente virtual amigable creada y desarrollada por Abdallah. "
            "Responde de forma útil, clara y natural en español."
        )
        
        response = co.chat(
            message=prompt_usuario,
            preamble=instrucciones
        )
        
        return response.text.strip()

    except Exception as e:
        print(f"Error Cohere: {e}")
        return f"Error al consultar la IA ({str(e)}). Revisa tu COHERE_API_KEY."

# --- LIMPIADOR DE MARKDOWN PARA PDF ---
def limpiar_markdown_pdf(texto):
    # Convierte encabezados (#, ##, ###) en texto en negrita
    texto = re.sub(r'#{1,6}\s*(.*)', r'<b>\1</b>', texto)
    # Convierte **texto** a <b>texto</b>
    texto = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', texto)
    # Elimina las líneas decorativas ---
    texto = re.sub(r'---', '', texto)
    return texto

# --- GENERADORES DE ARCHIVOS DINÁMICOS ---
def generar_pdf(nombre_archivo, titulo, contenido):
    doc = SimpleDocTemplate(nombre_archivo, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Limpiar formato Markdown sobrante
    contenido_procesado = limpiar_markdown_pdf(contenido)
    texto_limpio = contenido_procesado.replace('\n', '<br/>')
    
    story = [
        Paragraph(f"<b>{titulo}</b>", styles['Heading1']),
        Spacer(1, 12),
        Paragraph(texto_limpio, styles['BodyText'])
    ]
    doc.build(story)

def generar_word_texto(nombre_archivo, titulo, texto):
    doc = Document()
    doc.add_heading(titulo, level=1)
    for parrafo in texto.split('\n'):
        if parrafo.strip():
            # Limpiar negritas de Markdown para Word
            parrafo_limpio = re.sub(r'\*\*(.*?)\*\*', r'\1', parrafo)
            parrafo_limpio = re.sub(r'#{1,6}\s*', '', parrafo_limpio)
            doc.add_paragraph(parrafo_limpio.strip())
    doc.save(nombre_archivo)

def generar_excel(nombre_archivo):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Area", "Estado", "Creado Por"])
    ws.append(["Educación Física", "Activo", "Abdallah"])
    wb.save(nombre_archivo)

# --- DESCARGAR MÚSICA ---
def descargar_musica(nombre_cancion):
    archivo_salida = "cancion.mp3"
    if os.path.exists(archivo_salida):
        os.remove(archivo_salida)

    ydl_opts = {
        'format': 'bestaudio/best',
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'mp3',
            'preferredquality': '192',
        }],
        'outtmpl': 'cancion',
        'default_search': 'ytsearch1:',
        'quiet': True
    }

    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([nombre_cancion])
    
    return archivo_salida

# --- COMANDOS Y BOTONES ---
@bot.message_handler(commands=['start', 'help', 'modo'])
def send_welcome(message):
    markup = telebot.types.ReplyKeyboardMarkup(row_width=2, resize_keyboard=True)
    btn_texto = telebot.types.KeyboardButton("💬 Modo Solo Texto")
    btn_voz = telebot.types.KeyboardButton("🎙️ Modo Texto + Voz")
    markup.add(btn_texto, btn_voz)
    
    bot.send_message(
        message.chat.id, 
        "¡Hola! Soy Sofía, creada por Abdallah.\n\nElige cómo quieres que te responda:",
        reply_markup=markup
    )

# 1. Fotos
@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    bot.send_chat_action(message.chat.id, 'typing')
    bot.reply_to(message, "¡Recibí tu imagen! Qué gran foto me has mandado.")

# 2. Videos (Extraer Audio)
@bot.message_handler(content_types=['video'])
def handle_video(message):
    bot.send_message(message.chat.id, "🎥 Recibí tu video. Extrayendo el audio, dame un momento...")
    bot.send_chat_action(message.chat.id, 'upload_document')
    
    try:
        file_info = bot.get_file(message.video.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        
        video_path = "video_recibido.mp4"
        audio_path = "audio_extraido.mp3"
        
        with open(video_path, 'wb') as new_file:
            new_file.write(downloaded_file)
            
        os.system(f"ffmpeg -i {video_path} -q:a 0 -map a {audio_path} -y")
        
        if os.path.exists(audio_path):
            with open(audio_path, 'rb') as audio_file:
                bot.send_audio(message.chat.id, audio_file, title="Audio extraído", performer="Sofía Bot")
        else:
            bot.reply_to(message, "No pude extraer el audio de este formato de video.")
            
        for f in [video_path, audio_path]:
            if os.path.exists(f):
                os.remove(f)
                
    except Exception as e:
        print(f"Error procesando video: {e}")
        bot.reply_to(message, "Ocurrió un error al procesar el video.")

# 3. Conversación General
@bot.message_handler(func=lambda message: True)
def handle_conversation(message):
    user_id = message.chat.id
    user_text = message.text

    if user_text == "💬 Modo Solo Texto":
        modo_usuario[user_id] = "texto"
        bot.send_message(user_id, "✅ Modo activado: Sofía responderá solo en **Texto**.")
        return
    elif user_text == "🎙️ Modo Texto + Voz":
        modo_usuario[user_id] = "voz"
        bot.send_message(user_id, "✅ Modo activado: Sofía responderá con **Texto y Nota de Voz**.")
        return

    user_text_lower = user_text.lower()
    
    # Crear PDF Inteligente
    if "pdf" in user_text_lower:
        bot.send_chat_action(user_id, 'upload_document')
        tema = user_text_lower.replace("crea un pdf de", "").replace("crea un pdf sobre", "").replace("haz un pdf", "").replace("pdf", "").strip()
        if not tema:
            tema = "Información General"
            
        prompt_ia = f"Escribe una guía completa, clara y detallada sobre: {tema}."
        contenido_ia = consultar_ia_gratis(prompt_ia)
        
        archivo = f"documento_{user_id}.pdf"
        generar_pdf(archivo, f"Documento sobre {tema.capitalize()}", contenido_ia)
        
        with open(archivo, "rb") as f:
            bot.send_document(user_id, f, caption=f"📄 PDF listo sobre: *{tema.capitalize()}*", parse_mode="Markdown")
        
        if os.path.exists(archivo):
            os.remove(archivo)
        return

    # Crear Word Inteligente
    if "word" in user_text_lower or "doc" in user_text_lower or "crucigrama" in user_text_lower:
        bot.send_chat_action(user_id, 'upload_document')
        tema = user_text_lower.replace("crea un word de", "").replace("crea un word sobre", "").replace("haz un word", "").replace("word", "").replace("doc", "").strip()
        if not tema:
            tema = "Documento Informativo"
            
        prompt_ia = f"Redacta un documento detallado, articulado y formal sobre: {tema}."
        contenido_ia = consultar_ia_gratis(prompt_ia)
        
        archivo = f"documento_{user_id}.docx"
        generar_word_texto(archivo, tema.capitalize(), contenido_ia)
        
        with open(archivo, "rb") as f:
            bot.send_document(user_id, f, caption=f"📝 Documento de Word listo: *{tema.capitalize()}*", parse_mode="Markdown")
            
        if os.path.exists(archivo):
            os.remove(archivo)
        return
        
    # Crear Excel
    if "excel" in user_text_lower or "tabla" in user_text_lower:
        bot.send_chat_action(user_id, 'upload_document')
        archivo = "tabla.xlsx"
        generar_excel(archivo)
        with open(archivo, "rb") as f:
            bot.send_document(user_id, f)
        return

    # Descargar Música
    if "descarga" in user_text_lower or "cancion" in user_text_lower or "busca la canción" in user_text_lower:
        bot.send_message(user_id, "🔍 Buscando y descargando tu música...")
        bot.send_chat_action(user_id, 'upload_document')
        
        try:
            busqueda = user_text_lower.replace("descarga", "").replace("busca la canción", "").replace("cancion", "").strip()
            archivo_audio = descargar_musica(busqueda)
            
            with open(archivo_audio, "rb") as audio:
                bot.send_audio(user_id, audio, title=busqueda.capitalize(), performer="Sofía Bot")
            
            if os.path.exists(archivo_audio):
                os.remove(archivo_audio)
            return
        except Exception as e:
            print(f"Error al descargar: {e}")
            bot.reply_to(message, "No pude descargar la canción. Intenta escribir el nombre exacto.")
            return

    # Crear Imágenes
    if "dibuja" in user_text_lower or "crea una imagen" in user_text_lower:
        bot.send_message(user_id, "🎨 Creando tu imagen...")
        bot.send_chat_action(user_id, 'upload_photo')
        
        try:
            prompt = user_text_lower.replace("crea una imagen de", "").replace("crea una imagen", "").replace("dibuja", "").strip()
            prompt_url = prompt.replace(" ", "%20")
            url_imagen = f"https://image.pollinations.ai/prompt/{prompt_url}?width=1024&height=1024&nologo=true"
            bot.send_photo(user_id, url_imagen, caption=f"Aquí tienes: {prompt.capitalize()}")
            return
        except Exception as e:
            print(f"Error al crear imagen: {e}")
            bot.reply_to(message, "No pude generar esa imagen, intenta con otra descripción.")
            return

    # Respuesta por IA
    bot.send_chat_action(user_id, 'typing')
    respuesta_texto = consultar_ia_gratis(user_text)
    
    bot.send_message(user_id, respuesta_texto)
    
    if modo_usuario.get(user_id) == "voz":
        try:
            bot.send_chat_action(user_id, 'record_audio')
            archivo_voz = f"voz_{user_id}.mp3"
            tts = gTTS(text=respuesta_texto, lang='es', tld='com')
            tts.save(archivo_voz)
            
            with open(archivo_voz, "rb") as voice:
                bot.send_voice(user_id, voice)
                
            if os.path.exists(archivo_voz):
                os.remove(archivo_voz)
        except Exception as e:
            print(f"Error generando audio: {e}")

if __name__ == "__main__":
    bot.infinity_polling()
        
